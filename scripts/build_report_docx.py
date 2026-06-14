"""Generate the plant disease classification research report as a Word document."""
import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

PROJECT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def set_korean_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    set_korean_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return h


def add_para(doc, text, bold=False, size=11, align=None, space_after=6, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                set_korean_font(r, bold=True, size=10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    set_korean_font(r, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_figure(doc, image_path, caption, width_cm=14):
    if os.path.isfile(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cap.add_run(caption)
        set_korean_font(cr, size=10, color=RGBColor(80, 80, 80))
        cap.paragraph_format.space_after = Pt(12)
    else:
        add_para(doc, f"[이미지 없음: {image_path}]", size=10, color=RGBColor(180, 0, 0))


def build_report():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("딥러닝 기반 농작물 잎 이미지 병해충 진단 및\n도메인 시프트 극복 연구 보고서")
    set_korean_font(tr, size=18, bold=True)
    add_para(doc, "", space_after=12)

    # ── 2. 데이터셋 및 전처리 ──
    add_heading(doc, "2. 데이터셋 및 전처리", 1)

    add_heading(doc, "2.1 데이터셋 구성", 2)
    add_para(
        doc,
        "본 연구에서는 Kaggle New Plant Diseases Dataset을 기반으로 모델을 학습하였다. "
        "해당 데이터셋은 PlantVillage 기반의 증강 데이터셋으로, 총 10종 작물과 38개 클래스의 "
        "정상 및 질병 이미지를 포함한다. 학습 이미지는 70,295장, 검증 이미지는 17,572장으로 "
        "구성되어 있으며, Apple, Tomato, Potato, Grape, Corn 등 다양한 작물의 질병 클래스가 포함되어 있다.",
    )
    add_para(
        doc,
        "예를 들어 Apple 클래스는 Apple Scab, Black Rot, Cedar Apple Rust, Healthy로 구성되어 있고, "
        "Tomato 클래스는 Bacterial Spot, Early Blight, Late Blight, Leaf Mold, Mosaic Virus, Healthy 등으로 "
        "세분화되어 있다. 따라서 본 연구의 과제는 단순히 병해충 여부를 판단하는 이진 분류가 아니라, "
        "작물 종류와 질병 종류를 함께 구분하는 38개 클래스 다중 분류 문제이다.",
    )
    add_para(
        doc,
        "또한 실제 현장 적용 가능성을 검증하기 위해 야외 환경에서 촬영된 PlantDoc 데이터셋을 별도로 "
        "활용하였다. PlantDoc은 흙, 잡초, 가지, 그림자 등 복잡한 배경이 포함된 in-the-wild 이미지로, "
        "학습용 2,589장(train)과 테스트용 251장(test)으로 구성된다. 본 데이터셋은 도메인 시프트 분석 및 "
        "파인튜닝 실험의 평가 기준으로 사용되었다.",
    )
    add_table(
        doc,
        ["구분", "내용"],
        [
            ["학습 데이터셋", "Kaggle New Plant Diseases Dataset (PlantVillage 기반)"],
            ["학습 이미지 수", "70,295장"],
            ["검증 이미지 수", "17,572장"],
            ["클래스 수", "38개"],
            ["포함 작물 예시", "Apple, Tomato, Potato, Grape, Corn 등"],
            ["분류 유형", "정상 잎 및 병해충 잎 다중 분류"],
            ["야외 평가 데이터셋", "PlantDoc (train 2,589장 / test 251장)"],
        ],
        col_widths=[4.5, 12],
    )

    add_heading(doc, "2.2 전처리 및 데이터 증강", 2)
    add_para(
        doc,
        "모든 입력 이미지는 모델의 입력 크기에 맞춰 224×224 해상도로 변환하였다. "
        "ImageNet 사전학습 모델과의 입력 분포를 맞추기 위해 ImageNet 평균과 표준편차 "
        "(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])를 기준으로 정규화를 수행하였다. "
        "이후 PyTorch 학습에 사용하기 위해 이미지를 Tensor 형식으로 변환하였다.",
    )
    add_para(
        doc,
        "실험실 데이터(PlantVillage) 학습에는 Albumentations 기반 데이터 증강을 적용하였다. "
        "적용한 증강 기법은 Rotate(limit=30), HorizontalFlip, ColorJitter이다. "
        "Rotate는 촬영 각도 변화에 대한 강건성을 확보하기 위해 사용하였고, HorizontalFlip은 "
        "잎의 좌우 방향 변화에 대응하기 위해 적용하였다. ColorJitter는 조명, 밝기, 대비, 채도 변화에 "
        "대한 모델의 적응력을 높이기 위해 사용하였다. 검증 데이터에는 증강을 적용하지 않고 "
        "Resize, Normalize, ToTensor만 적용하여 평가 결과가 왜곡되지 않도록 하였다.",
    )
    add_para(
        doc,
        "야외 데이터(PlantDoc) 파인튜닝 단계에서는 보다 강한 증강을 적용하였다. "
        "RandomResizedCrop(scale=0.7~1.0), RandomHorizontalFlip, RandomRotation(30°), "
        "ColorJitter(brightness/contrast/saturation=0.3)를 사용하여 실제 현장의 조명·크롭·각도 변화를 "
        "모방하였다.",
    )

    # ── 3. 모델 구조 및 학습 방법 ──
    add_heading(doc, "3. 모델 구조 및 학습 방법", 1)

    add_heading(doc, "3.1 베이스라인 CNN", 2)
    add_para(
        doc,
        "초기 모델로는 4개의 Convolution Block으로 구성된 자체 CNN 모델을 설계하였다. "
        "각 블록은 Conv2d, BatchNorm2d, ReLU, MaxPool2d로 구성되며, 채널 수는 32, 64, 128, 256으로 "
        "점진적으로 증가한다. 입력 이미지는 3×224×224 크기이며, 네 개의 블록을 통과한 뒤 "
        "256×14×14 크기의 feature map으로 변환된다.",
    )
    add_para(
        doc,
        "초기 베이스라인 모델은 Flatten 계층을 사용하여 feature map을 1차원 벡터로 펼친 뒤, "
        "Fully Connected Layer를 통해 최종적으로 38개 클래스를 분류하였다. "
        "Baseline CNN은 약 2,610만 개의 파라미터를 가지며, 학습 비용과 과적합 가능성을 높이는 요인으로 작용할 수 있다.",
    )

    add_heading(doc, "3.2 모델 경량화 및 전이학습 모델 비교", 2)
    add_para(
        doc,
        "Baseline CNN의 파라미터 문제를 해결하기 위해 Global Average Pooling(GAP)을 도입한 경량화 모델을 실험하였다. "
        "GAP를 적용하면 파라미터 수가 약 2,610만 개에서 약 54만 개로 크게 감소하였다. "
        "이후 전이학습 기반 모델로 ResNet50, MobileNetV3, EfficientNet-B0를 비교하였다.",
    )
    add_table(
        doc,
        ["모델", "주요 구조", "파라미터 수", "특징"],
        [
            ["Baseline CNN", "Flatten + Dense", "약 26.1M", "가장 무거운 구조"],
            ["Baseline CNN + GAP", "GAP + Dense", "약 0.5M", "경량화 효과 큼"],
            ["ResNet50", "GAP + Residual Connection", "약 23.0M", "깊은 층의 안정적 학습"],
            ["MobileNetV3", "GAP + Depthwise Conv", "약 4.2M", "모바일 환경에 적합"],
            ["EfficientNet-B0", "GAP + Compound Scaling", "약 4.0M", "성능과 효율의 균형 우수"],
            ["ViT-Small", "Self-Attention + Patch Embedding", "약 21.7M", "글로벌 문맥 활용 (후속 실험)"],
        ],
        col_widths=[3.5, 4.5, 2.5, 5.5],
    )
    add_para(
        doc,
        "실험실 데이터 학습 단계에서는 EfficientNet-B0가 정확도와 파라미터 효율성 측면에서 "
        "가장 우수한 결과를 보였기 때문에, 초기 도메인 시프트 분석 및 가설 검증의 중심 모델로 사용하였다.",
    )

    add_heading(doc, "3.3 학습 설정", 2)
    add_para(
        doc,
        "모델 학습은 PyTorch 2.x와 CUDA 기반 GPU 환경에서 수행하였다. "
        "손실 함수는 CrossEntropyLoss를 사용하였으며, 옵티마이저는 Adam을 사용하였다. "
        "학습률은 0.001, 배치 사이즈는 32, 에포크 수는 10으로 설정하였다. "
        "과적합 방지를 위해 Dropout 0.5를 적용하였다. "
        "검증 정확도가 최고치를 갱신할 때마다 best_model.pth로 가중치를 저장하였다.",
    )
    add_para(
        doc,
        "야외 데이터 파인튜닝(가설 3 및 ViT 비교 실험)에서는 Label Smoothing(0.1), "
        "Cosine Annealing Learning Rate, 2단계 학습(Phase 1: 분류 헤드 5 epoch → Phase 2: 전체 25 epoch)을 "
        "공통으로 적용하여 공정한 비교가 이루어지도록 하였다.",
    )

    # ── 4. 실험 결과 및 도메인 시프트 분석 ──
    add_heading(doc, "4. 실험 결과 및 도메인 시프트 분석", 1)

    add_heading(doc, "4.1 실험실 데이터 성능", 2)
    add_para(
        doc,
        "EfficientNet-B0 모델을 PlantVillage 기반 학습 데이터로 10 Epoch 동안 학습한 결과, "
        "검증 데이터에서 최고 정확도 99.73%를 달성하였다. Best Validation Loss는 0.0094였으며, "
        "최고 성능은 7번째 Epoch에서 나타났다. 이는 실험실 환경의 정제된 잎 이미지에 대해서는 "
        "모델이 38개 클래스를 매우 높은 정확도로 분류할 수 있음을 의미한다.",
    )
    add_figure(
        doc,
        os.path.join(PROJECT_DIR, "training_results_efficientnet_b0.png"),
        "그림 1. EfficientNet-B0 학습 결과 (Best Val Acc: 99.73%, Best Epoch: 7, Best Val Loss: 0.0094)",
    )
    add_para(
        doc,
        "그러나 이러한 결과는 실험실 환경에 한정된 성능일 수 있다. PlantVillage 데이터는 대부분 "
        "단색 배경 위에 잎 하나가 명확히 촬영되어 있으며, 실제 농업 현장의 복잡한 배경과는 차이가 크다. "
        "따라서 모델의 실제 적용 가능성을 검증하기 위해 야외 데이터셋인 PlantDoc 이미지에 대해 "
        "추가 테스트를 수행하였다.",
    )

    add_heading(doc, "4.2 야외 데이터 테스트 결과", 2)
    add_para(
        doc,
        "실험실 데이터에서 99.73%의 정확도를 보인 EfficientNet-B0 모델을 야외 데이터(PlantDoc test 251장)에 "
        "적용한 결과, 정확도는 19.12%로 급격히 하락하였다. 야외 데이터는 흙, 잡초, 줄기, 가지, 그림자, "
        "다양한 조명 조건 등이 함께 포함되어 있어, 모델이 병변의 본질적 특징보다 실험실 데이터의 "
        "배경 패턴에 과적합되었을 가능성이 높다.",
    )
    add_table(
        doc,
        ["평가 환경", "데이터 특성", "정확도"],
        [
            ["실험실 데이터 (PlantVillage Val)", "단색 배경, 정제된 잎 이미지", "99.73%"],
            ["야외 데이터 (PlantDoc Test)", "흙, 잡초, 그림자 등 복잡한 배경", "19.12%"],
        ],
        col_widths=[5, 7, 3],
    )

    add_heading(doc, "4.3 Grad-CAM 기반 원인 분석", 2)
    add_para(
        doc,
        "도메인 시프트의 원인을 구체적으로 분석하기 위해 Grad-CAM을 활용하였다. "
        "Grad-CAM은 모델이 예측을 수행할 때 이미지의 어느 영역에 집중했는지 히트맵 형태로 시각화하는 "
        "설명 가능한 인공지능(XAI) 기법이다.",
    )
    add_para(
        doc,
        "분석 결과, 실험실에서 학습된 EfficientNet-B0 모델은 잎의 병변 영역이 아니라 흙, 가지, 배경 영역에 "
        "강하게 반응하는 경우가 확인되었다. 예를 들어 실제 정답이 Apple Scab인 이미지에 대해 모델이 "
        "Peach Bacterial Spot으로 오진한 사례에서, 히트맵은 잎 내부의 병변보다 배경에 더 강하게 활성화되어 있었다.",
    )
    add_figure(
        doc,
        os.path.join(PROJECT_DIR, "gradcam_results", "wrong_2_true_Apple___Apple_scab_pred_Peach___Bacterial_spot.png"),
        "그림 2. Grad-CAM 시각화 — 정답: Apple Scab / 예측: Peach Bacterial Spot (배경 영역 활성화)",
        width_cm=12,
    )

    # ── 5. 도메인 시프트 극복을 위한 가설 검증 ──
    add_heading(doc, "5. 도메인 시프트 극복을 위한 가설 검증", 1)

    add_heading(doc, "5.1 가설 1: 물리적 전처리 기반 배경 제거", 2)
    add_para(
        doc,
        "첫 번째 가설은 배경을 제거함으로써 잎의 병변에 집중하도록 만들 수 있다는 것이다. "
        "rembg 기반 AI 배경 제거를 적용한 뒤 파인튜닝한 결과, 정확도는 47.81%로 향상되었으나 "
        "잎 가장자리·병변 일부가 훼손되는 문제가 발생하여 충분한 성능을 달성하지 못하였다.",
    )

    add_heading(doc, "5.2 가설 2: 실험실 데이터와 야외 데이터의 혼합 학습", 2)
    add_para(
        doc,
        "두 번째 가설은 실험실 데이터 5,000장과 야외 데이터 2,500장을 혼합하여 파인튜닝하는 것이다. "
        "실험 결과 정확도는 47.01%로, 가설 1과 유사한 수준이며 기대만큼의 성능 향상을 보이지 못하였다. "
        "쉬운 실험실 데이터가 학습 방향을 분산시켜 도메인 적응 효과를 약화시킨 것으로 해석된다.",
    )

    add_heading(doc, "5.3 가설 3: 야외 데이터 중심의 본질적 도메인 적응", 2)
    add_para(
        doc,
        "세 번째 가설은 PlantDoc 야외 데이터만을 중심으로 전이학습을 수행하고, "
        "강한 데이터 증강(ColorJitter, RandomResizedCrop)과 Cosine Annealing LR, Label Smoothing을 "
        "함께 적용하는 것이다. EfficientNet-B0 기준 실험 결과 정확도는 60.56%로 상승하였으며, "
        "기존 19.12% 대비 3배 이상 향상된 결과이다. 소량(2,589장)의 야외 데이터만으로도 "
        "전이학습과 강한 증강을 결합하면 유의미한 도메인 적응이 가능함을 확인하였다.",
    )
    add_figure(
        doc,
        os.path.join(PROJECT_DIR, "assets", "curve_v2.png"),
        "그림 3. 가설 3(EfficientNet V2) 학습 곡선 — 최종 야외 정확도 60.56%",
        width_cm=12,
    )
    add_table(
        doc,
        ["구분", "전략", "정확도", "결과"],
        [
            ["기존 모델", "실험실 가중치 그대로 야외 테스트", "19.12%", "실패"],
            ["가설 1", "rembg 기반 배경 제거 후 파인튜닝", "47.81%", "제한적 개선"],
            ["가설 2", "실험실+야외 데이터 혼합 학습", "47.01%", "제한적 개선"],
            ["가설 3", "야외 데이터 중심 증강 및 최적화 전이학습", "60.56%", "CNN 기준 최선"],
        ],
        col_widths=[2.5, 6, 2.5, 3],
    )

    # ── 6. Vision Transformer 비교 실험 ──
    add_heading(doc, "6. Vision Transformer(ViT) 아키텍처 비교 실험", 1)
    add_para(
        doc,
        "가설 3에서 EfficientNet-B0가 도메인 적응의 유효성을 입증한 이후, Self-Attention 기반의 "
        "Vision Transformer(ViT-Small, vit_small_patch16_224)가 야외 환경에서 더 우수한 성능을 "
        "보이는지 검증하였다. 공정한 비교를 위해 EfficientNet-B0를 동일 프로토콜로 재학습하고 "
        "ViT-Small과 head-to-head 비교를 수행하였다.",
    )
    add_para(
        doc,
        "비교 실험의 공통 조건은 다음과 같다. (1) 동일 데이터 분할: PlantDoc train/test, "
        "(2) 동일 증강: RandomResizedCrop, ColorJitter, RandomRotation, "
        "(3) 동일 손실·스케줄러: Label Smoothing(0.1) + Cosine Annealing, "
        "(4) 동일 2단계 학습: Phase 1 헤드 5 epoch → Phase 2 전체 25 epoch.",
    )
    add_table(
        doc,
        ["모델", "파라미터", "PlantDoc Test 정확도", "비고"],
        [
            ["EfficientNet-B0 (동일 프로토콜)", "약 405만", "65.34% (164/251)", "CNN 기준선"],
            ["ViT-Small", "약 2,168만", "70.92% (178/251)", "+5.58%p 우위"],
        ],
        col_widths=[4.5, 2.5, 4, 4],
    )
    add_para(
        doc,
        "Head-to-head 분석 결과, 251장의 테스트 이미지 중 149장은 두 모델 모두 정답, "
        "29장은 ViT만 정답, 15장은 EfficientNet만 정답이었다. ViT의 우위는 특히 Apple 계열 "
        "(89.7% vs 75.9%)과 같이 야외 복잡 배경에서 두드러졌다. Self-Attention이 이미지 전역의 "
        "문맥을 활용하여 배경 노이즈와 병변 영역을 보다 효과적으로 구분할 수 있기 때문으로 해석된다.",
    )

    add_heading(doc, "6.1 ViT Attention 시각화 분석", 2)
    add_para(
        doc,
        "ViT-Small 파인튜닝 모델에 대해 CLS 토큰의 attention map을 시각화하여, "
        "Grad-CAM 분석과 비교하였다. 마지막 Transformer 블록에서 CLS 토큰이 각 패치에 "
        "부여하는 attention 가중치(멀티헤드 평균)를 히트맵으로 오버레이하였다.",
    )
    add_para(
        doc,
        "정답 사례(Apple Scab)에서는 잎의 검은 반점(병변)에 attention이 집중되는 패턴이 관찰되어, "
        "파인튜닝 후 모델이 병변 특징을 활용하고 있음을 확인하였다. 반면 일부 정답·오답 사례에서는 "
        "이미지 모서리나 배경에 attention이 쏠리는 shortcut 패턴도 잔존하였다.",
    )
    add_para(
        doc,
        "오답 사례는 크게 세 유형으로 분류된다. "
        "(1) 유사 병해 혼동: Apple Scab → Cedar Apple Rust — 병변은 인식하나 클래스 구분 실패. "
        "(2) 병해 미검출: Cedar Apple Rust → Healthy — 배경 그림자에 attention 집중. "
        "(3) 종(species) 혼동: Blueberry Healthy → Apple/Soybean Healthy — 잎 전체 형태보다 "
        "배경·국소 texture에 의존.",
    )

    add_figure(
        doc,
        os.path.join(PROJECT_DIR, "vit_attention_results", "correct_1_true_Apple___Apple_scab_pred_Apple___Apple_scab.png"),
        "그림 4. ViT Attention — 정답 사례 (Apple Scab): 병변 영역에 attention 집중",
        width_cm=11,
    )
    add_figure(
        doc,
        os.path.join(PROJECT_DIR, "vit_attention_results", "wrong_2_true_Apple___Cedar_apple_rust_pred_Apple___healthy.png"),
        "그림 5. ViT Attention — 오답 사례 (Cedar Apple Rust → Healthy): 배경 그림자에 attention 집중",
        width_cm=11,
    )

    add_table(
        doc,
        ["비교 항목", "Grad-CAM (실험실 EfficientNet)", "ViT Attention (야외 적응 ViT)"],
        [
            ["분석 대상 모델", "PlantVillage 학습 EfficientNet-B0", "PlantDoc 파인튜닝 ViT-Small"],
            ["주요 attention 경향", "흙·가지·배경에 집중", "병변 집중 사례 증가"],
            ["잔존 문제", "배경 shortcut이 지배적", "일부 케이스에서 배경 shortcut 잔존"],
            ["대응 정확도", "야외 19.12%", "야외 70.92%"],
        ],
        col_widths=[4, 5.5, 5.5],
    )

    # ── 7. 종합 비교 및 결론 ──
    add_heading(doc, "7. 종합 비교 및 결론", 1)
    add_para(
        doc,
        "본 연구에서는 딥러닝 기반 농작물 잎 이미지 병해충 진단 시스템을 구축하고, "
        "실험실 데이터와 야외 데이터 사이의 도메인 시프트 문제를 분석하였다. "
        "EfficientNet-B0는 PlantVillage 검증 데이터에서 99.73%의 높은 정확도를 달성하였으나, "
        "야외 데이터에서는 19.12%로 크게 하락하였다. Grad-CAM 분석을 통해 모델이 병변보다 "
        "배경에 집중하는 경향을 확인하였다.",
    )
    add_para(
        doc,
        "세 가지 가설 검증 결과, 배경 제거(47.81%), 데이터 혼합(47.01%)보다 "
        "야외 데이터 중심 파인튜닝(60.56%)이 가장 효과적이었다. "
        "이후 동일 프로토콜에서 ViT-Small을 도입한 결과 70.92%로 현재 최고 성능을 달성하였으며, "
        "EfficientNet(65.34%) 대비 5.58%p 향상을 기록하였다. ViT Attention 시각화는 "
        "파인튜닝 후 병변 집중이 개선되었음을 보여주었으나, 유사 병해 혼동과 healthy 종 혼동 등 "
        "잔존 오류 유형도 확인하였다.",
    )
    add_table(
        doc,
        ["단계", "모델/전략", "야외 Test 정확도"],
        [
            ["실험실 학습", "EfficientNet-B0 (PlantVillage)", "99.73% (Val)"],
            ["도메인 시프트", "실험실 가중치 → PlantDoc", "19.12%"],
            ["가설 3", "EfficientNet V2 파인튜닝", "60.56%"],
            ["아키텍처 비교", "EfficientNet (동일 프로토콜)", "65.34%"],
            ["아키텍처 비교", "ViT-Small (동일 프로토콜)", "70.92%"],
        ],
        col_widths=[3.5, 6.5, 4],
    )
    add_para(
        doc,
        "결론적으로, 실제 적용 환경과 학습 환경의 차이가 큰 경우 목표 도메인에 집중한 파인튜닝이 "
        "필수적이며, 아키텍처 측면에서는 Self-Attention 기반 ViT가 CNN 대비 추가적인 성능 향상을 "
        "제공할 수 있음을 확인하였다.",
    )

    add_heading(doc, "7.1 향후 연구 방향", 2)
    add_para(
        doc,
        "향후 연구에서는 현재 최고 성능 모델인 ViT-Small(70.92%)을 기반으로 다음 방향을 추진할 수 있다. "
        "첫째, DANN(적대적 도메인 적응)을 적용하여 실험실·야외 데이터 간 분포 차이를 추가로 완화한다. "
        "둘째, YOLO 기반 객체 탐지로 잎 영역을 먼저 검출한 뒤 분류하는 파이프라인을 구성하여 "
        "Attention 분석에서 확인된 배경 shortcut을 줄인다. 셋째, 야외 식물 이미지에 대한 "
        "자기지도학습(Pre-training) 후 파인튜닝으로 라벨 부족 상황에서도 강건한 특징 표현을 학습한다. "
        "목표는 야외 정확도 80~90% 달성이다.",
    )
    add_para(
        doc,
        "본 연구는 딥러닝 모델 평가에서 실험실 정확도뿐 아니라 실제 적용 환경에서의 일반화 성능 검증이 "
        "필수적임을 보여준다. 농업 인공지능 시스템이 현장에서 활용되기 위해서는 도메인 차이를 인식하고, "
        "도메인 적응 전략과 적합한 아키텍처 선택을 함께 고려해야 한다.",
    )

    out_paths = [
        os.path.join(PROJECT_DIR, "연구보고서_농작물병해충진단.docx"),
        os.path.join(PROJECT_DIR, "research_report_plant_disease.docx"),
    ]
    saved = []
    for out_path in out_paths:
        try:
            doc.save(out_path)
            saved.append(out_path)
            print(f"Saved: {out_path}")
        except PermissionError:
            alt = out_path.replace(".docx", "_v2.docx")
            doc.save(alt)
            saved.append(alt)
            print(f"Saved (alt): {alt}  (original file may be open)")
    return saved[0] if saved else None


if __name__ == "__main__":
    build_report()
